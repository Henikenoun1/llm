"""
Generate a Word guide describing the exact data files and templates required
to prepare a new domain for training and inference.
"""
from __future__ import annotations

from pathlib import Path


ROOT = Path(__file__).resolve().parents[1]
TEMPLATES_DIR = ROOT / "data" / "templates"
DOCS_DIR = ROOT / "docs"
DOCS_DIR.mkdir(parents=True, exist_ok=True)


def _read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8").strip()


def _add_code_block(document, text: str) -> None:
    for line in text.splitlines():
        paragraph = document.add_paragraph()
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        if hasattr(run.font, "size"):
            from docx.shared import Pt

            run.font.size = Pt(9)


def generate_data_guide() -> Path:
    try:
        from docx import Document
    except Exception as exc:  # pragma: no cover - dependency installed in runtime
        raise RuntimeError("python-docx is required to generate the Word guide.") from exc

    document = Document()
    document.add_heading("Call Center LLM Data Preparation Guide", level=0)
    document.add_paragraph(
        "This guide explains every data file, template, and prerequisite needed to prepare a new domain, "
        "train the model, evaluate it, and operate it in production."
    )

    document.add_heading("1. Global Pipeline", level=1)
    for item in [
        "1. Define the business domain in data/config/domain.json.",
        "2. Add few-shot examples in data/config/few_shots.jsonl.",
        "3. Prepare self-supervised language texts.",
        "4. Prepare SFT business dialogues.",
        "5. Prepare DPO preference pairs.",
        "6. Add knowledge base files in data/kb and reference documents in data/rag.",
        "7. Add inference evaluation cases in data/eval/inference_cases.jsonl.",
        "8. Run download -> prepare -> self-sup -> sft -> dpo -> eval -> serve.",
    ]:
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("2. Domain Configuration File", level=1)
    document.add_paragraph(
        "Main file: data/config/domain.json. This file defines the business, intents, slots, tools, runtime mode, "
        "RAG paths, memory paths, API settings, and database settings."
    )
    document.add_paragraph(
        "Dataset entries should contain either a Hugging Face dataset identifier in hf_dataset, or a local_path. "
        "The dataset name alone is not enough unless it is a valid Hugging Face repository id. "
        "If you use an external link, download it first and point local_path to the local file."
    )
    _add_code_block(document, _read_text(TEMPLATES_DIR / "domain.template.json"))

    document.add_heading("3. Few-Shot File", level=1)
    document.add_paragraph(
        "Main file: data/config/few_shots.jsonl. Each line is one short example with an intent, a user message, "
        "and a preferred assistant answer. Few-shots are used at inference time to reinforce style and routing."
    )
    _add_code_block(document, _read_text(TEMPLATES_DIR / "few_shots.template.jsonl"))

    document.add_heading("4. Self-Supervised Dataset", level=1)
    document.add_paragraph(
        "Purpose: teach the model the language, expressions, tone, spelling variation, and customer phrasing. "
        "Format: one JSONL line with a single text field."
    )
    document.add_paragraph(
        "Good data for this file: raw customer language, support language, dialectal speech, transcribed text, "
        "and large quantities of clean target-language content. This step is for language adaptation, not business policy."
    )
    _add_code_block(document, _read_text(TEMPLATES_DIR / "self_supervised.template.jsonl"))

    document.add_heading("5. SFT Dialogue Dataset", level=1)
    document.add_paragraph(
        "Purpose: teach the business task. Format: one JSONL line with a messages array. "
        "This is the most important dataset for call-center behavior."
    )
    document.add_paragraph(
        "Each conversation should reflect the real production style: asking for missing fields, confirming actions, "
        "tracking orders, handling objections, and speaking in the exact tone you want in production."
    )
    _add_code_block(document, _read_text(TEMPLATES_DIR / "sft_dialogues.template.jsonl"))

    document.add_heading("6. DPO Preference Dataset", level=1)
    document.add_paragraph(
        "Purpose: teach response preference. Format: prompt, chosen, rejected. "
        "Chosen must be the better answer. Rejected must be plausible but worse, unsafe, incomplete, or stylistically bad."
    )
    _add_code_block(document, _read_text(TEMPLATES_DIR / "dpo_pairs.template.jsonl"))

    document.add_heading("7. RAG Documents", level=1)
    document.add_paragraph(
        "Purpose: external business knowledge used at inference time. Store policies, FAQs, procedures, delivery rules, "
        "pricing references, escalation rules, scripts, and internal knowledge in data/rag or data/kb."
    )
    document.add_paragraph(
        "Supported files include .md, .txt, .jsonl, .json, and .csv. Documents are chunked automatically and embedded for retrieval."
    )
    _add_code_block(document, _read_text(TEMPLATES_DIR / "rag_document.template.md"))

    document.add_heading("8. Tool and Knowledge Base Files", level=1)
    document.add_paragraph(
        "Tools are configured in code, but the business data they use should live in files under data/kb. "
        "Typical files are product catalogues, store schedules, policy documents, and mock order exports."
    )
    document.add_paragraph("Product catalogue example:")
    _add_code_block(document, _read_text(TEMPLATES_DIR / "kb_products.template.csv"))
    document.add_paragraph("Store information example:")
    _add_code_block(document, _read_text(TEMPLATES_DIR / "kb_stores.template.csv"))

    document.add_heading("9. Inference Evaluation Cases", level=1)
    document.add_paragraph(
        "Main file: data/eval/inference_cases.jsonl. This file is used to score the production inference pipeline. "
        "Each line may contain the input, expected intent, expected slots, expected tool, human review expectation, "
        "reference responses, required keywords, and forbidden keywords."
    )
    _add_code_block(document, _read_text(TEMPLATES_DIR / "inference_eval_cases.template.jsonl"))

    document.add_heading("10. Runtime Feedback and Learning Files", level=1)
    for item in [
        "data/history/conversations.jsonl: sanitized conversation log.",
        "data/history/learning_pending.jsonl: candidate examples generated by runtime.",
        "data/history/learning_buffer.jsonl: approved examples reused later for SFT.",
        "data/history/feedback_dpo.jsonl: approved preference corrections reused later for DPO.",
        "data/history/ratings_log.jsonl: light good/bad ratings.",
        "data/history/admin_corrections.jsonl: live admin corrections applied immediately at runtime.",
    ]:
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("11. Runtime Modes", level=1)
    for item in [
        "collect_execute: collect fields, structure the action, and keep responses operational.",
        "speak: natural conversation with controlled execution.",
        "autonomous: natural conversation with maximum autonomy based on configured tool policy.",
    ]:
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("12. Database and Docker", level=1)
    document.add_paragraph(
        "The application can run without a database, but for serious production usage you should enable PostgreSQL. "
        "The Docker stack provided in docker-compose.yml starts PostgreSQL and Adminer."
    )
    document.add_paragraph(
        "Use CALL_CENTER_DATABASE_URL to connect the application. Example: "
        "postgresql+psycopg://callcenter:callcenter@localhost:5432/callcenter"
    )
    document.add_paragraph(
        "The database mirrors session states, conversations, learning examples, ratings, feedback events, and live admin corrections."
    )

    document.add_heading("13. Recommended Minimum Data Before Training", level=1)
    for item in [
        "Self-supervised texts: as much clean target-language text as possible.",
        "SFT dialogues: at least hundreds of high-quality business conversations, ideally thousands.",
        "DPO pairs: hundreds of strong chosen/rejected pairs for the most important scenarios.",
        "Few-shots: 1 to 5 strong examples per critical intent.",
        "RAG documents: all policies, scripts, FAQs, pricing references, and operational procedures.",
        "Evaluation cases: at least 10 to 20 cases per major intent before production promotion.",
    ]:
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("14. Promotion Checklist", level=1)
    for item in [
        "Data preparation completed without malformed lines.",
        "Self-supervised, SFT, and DPO stages completed successfully.",
        "Evaluation report generated with acceptable intent accuracy, slot F1, tool accuracy, and response quality.",
        "Critical inference cases validated manually in the target runtime mode.",
        "Database enabled for production deployments.",
        "Admin correction workflow tested.",
    ]:
        document.add_paragraph(item, style="List Bullet")

    output_path = DOCS_DIR / "Call_Center_LLM_Data_Guide.docx"
    document.save(output_path)

    markdown_path = DOCS_DIR / "Call_Center_LLM_Data_Guide.md"
    markdown_path.write_text(
        "\n".join(
            [
                "# Call Center LLM Data Preparation Guide",
                "",
                "See the Word document for the full formatted version.",
                "",
                "Templates:",
                "- data/templates/domain.template.json",
                "- data/templates/few_shots.template.jsonl",
                "- data/templates/self_supervised.template.jsonl",
                "- data/templates/sft_dialogues.template.jsonl",
                "- data/templates/dpo_pairs.template.jsonl",
                "- data/templates/inference_eval_cases.template.jsonl",
                "- data/templates/rag_document.template.md",
                "- data/templates/kb_products.template.csv",
                "- data/templates/kb_stores.template.csv",
            ]
        ),
        encoding="utf-8",
    )

    return output_path


if __name__ == "__main__":
    print(generate_data_guide())
